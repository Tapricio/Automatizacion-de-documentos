import re
import pandas as pd
from docx import Document
from tkinter import Tk, filedialog
import os
import win32com.client

# ---------------------------------
# Inicio
# ---------------------------------
print("\n==============================")
print(" GENERADOR MASIVO DE DOCUMENTOS ")
print("==============================\n")

# ---------------------------------
# Selección archivos
# ---------------------------------
root = Tk()
root.withdraw()
root.attributes('-topmost', True)

print("1) Selecciona el archivo Excel...\n")

excel_path = filedialog.askopenfilename(
    title="Selecciona Excel",
    filetypes=[("Excel", "*.xlsx")]
)

print("2) Selecciona la plantilla Word...\n")

template_path = filedialog.askopenfilename(
    title="Selecciona Word",
    filetypes=[("Word", "*.docx")]
)

print("3) Selecciona la carpeta donde se guardarán los documentos...\n")

output_folder = filedialog.askdirectory(
    title="Selecciona carpeta destino"
)

nombre_base = input(
    "\n4) Escribe el nombre de la carpeta donde guardarás los documentos:\n\n> "
).strip() or "documentos generados"

filaPrint = 1

# ---------------------------------
# Validaciones
# ---------------------------------
if not excel_path:
    print("No seleccionaste Excel")
    exit()

if not template_path:
    print("No seleccionaste Word")
    exit()

if not output_folder:
    print("No seleccionaste carpeta")
    exit()

# ---------------------------------
# Crear carpeta única
# ---------------------------------
print("\nCreando carpeta de documentos...\n")

carpeta_generados = os.path.join(
    output_folder,
    nombre_base
)

contador = 1

while os.path.exists(carpeta_generados):

    carpeta_generados = os.path.join(
        output_folder,
        f"{nombre_base} ({contador})"
    )

    contador += 1

os.makedirs(carpeta_generados)

# ---------------------------------
# Leer Excel
# ---------------------------------
print("Leyendo Excel...\n")

df = pd.read_excel(excel_path)

# ---------------------------------
# Crear columna Adjuntos como texto
# ---------------------------------
if "Adjuntos" not in df.columns:
    df["Adjuntos"] = ""

df["Adjuntos"] = df["Adjuntos"].astype(str)

# ---------------------------------
# Columnas plantilla
# ---------------------------------
columnas_plantilla = [
    c for c in df.columns
    if c != "Adjuntos"
]

# ---------------------------------
# Mostrar columnas
# ---------------------------------
print("\nColumnas disponibles:\n")

for i, columna in enumerate(columnas_plantilla):
    print(f"{i + 1}. {columna}")

# ---------------------------------
# Elegir columnas nombre archivo
# ---------------------------------
seleccion = input(
    "\nEscribe números separados por coma para el nombre del archivo.\nEjemplo: 1,2\n\n> "
)

indices = [
    int(x.strip()) - 1
    for x in seleccion.split(",")
]

columnas_nombre = [
    columnas_plantilla[i]
    for i in indices
]

# ---------------------------------
# Texto adicional opcional
# ---------------------------------
texto_adicional = input(
    "\nTexto adicional opcional para el nombre del archivo (ENTER para omitir):\n\n> "
).strip()

print("\nNombre archivo usará:")
print(columnas_nombre)

if texto_adicional:
    print(f"Texto adicional: {texto_adicional}")

# ---------------------------------
# Funciones formato
# ---------------------------------
def obtener_formato_por_caracter(parrafo):

    formato_por_caracter = []

    for run in parrafo.runs:

        formato = {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "style": run.style,
            "font_name": run.font.name,
            "font_size": run.font.size,
            "highlight_color": run.font.highlight_color,
            "color_rgb": run.font.color.rgb if run.font.color else None,
            "strike": run.font.strike,
            "subscript": run.font.subscript,
            "superscript": run.font.superscript,
        }

        formato_por_caracter.extend(
            [formato] * len(run.text)
        )

    return formato_por_caracter


def aplicar_formato_run(run, formato):

    run.bold = formato["bold"]
    run.italic = formato["italic"]
    run.underline = formato["underline"]

    if formato["style"] is not None:
        run.style = formato["style"]

    if formato["font_name"] is not None:
        run.font.name = formato["font_name"]

    if formato["font_size"] is not None:
        run.font.size = formato["font_size"]

    run.font.highlight_color = formato["highlight_color"]
    run.font.strike = formato["strike"]
    run.font.subscript = formato["subscript"]
    run.font.superscript = formato["superscript"]

    if formato["color_rgb"] is not None:
        run.font.color.rgb = formato["color_rgb"]


# ---------------------------------
# Reemplazo conservando formato
# ---------------------------------
def reemplazar_en_parrafo(parrafo, reemplazos):

    texto_original = "".join(
        run.text for run in parrafo.runs
    )

    if not texto_original:
        return

    claves_ordenadas = sorted(
        reemplazos.keys(),
        key=lambda x: len(str(x)),
        reverse=True
    )

    patron = re.compile(
        "|".join(
            re.escape(str(k))
            for k in claves_ordenadas
        )
    )

    coincidencias = list(
        patron.finditer(texto_original)
    )

    if not coincidencias:
        return

    formato_por_caracter = obtener_formato_por_caracter(
        parrafo
    )

    segmentos = []

    ultima_pos = 0

    for coincidencia in coincidencias:

        inicio = coincidencia.start()
        fin = coincidencia.end()

        if inicio > ultima_pos:

            segmentos.append(
                (
                    texto_original[ultima_pos:inicio],
                    formato_por_caracter[ultima_pos:inicio]
                )
            )

        valor_reemplazo = str(
            reemplazos[
                coincidencia.group(0)
            ]
        )

        formato_reemplazo = (
            formato_por_caracter[inicio]
            if inicio < len(formato_por_caracter)
            else formato_por_caracter[-1]
        )

        segmentos.append(
            (
                valor_reemplazo,
                formato_reemplazo
            )
        )

        ultima_pos = fin

    if ultima_pos < len(texto_original):

        segmentos.append(
            (
                texto_original[ultima_pos:],
                formato_por_caracter[ultima_pos:]
            )
        )

    for run in list(parrafo.runs):
        run._element.getparent().remove(
            run._element
        )

    for texto_segmento, formato_segmento in segmentos:

        if not texto_segmento:
            continue

        if isinstance(formato_segmento, dict):

            run = parrafo.add_run(
                texto_segmento
            )

            aplicar_formato_run(
                run,
                formato_segmento
            )

            continue

        acumulado = ""
        ultimo_formato = formato_segmento[0]

        for caracter, formato_caracter in zip(
            texto_segmento,
            formato_segmento
        ):

            if (
                formato_caracter != ultimo_formato
                and acumulado
            ):

                run = parrafo.add_run(
                    acumulado
                )

                aplicar_formato_run(
                    run,
                    ultimo_formato
                )

                acumulado = caracter
                ultimo_formato = formato_caracter

            else:
                acumulado += caracter

        if acumulado:

            run = parrafo.add_run(
                acumulado
            )

            aplicar_formato_run(
                run,
                ultimo_formato
            )

# ---------------------------------
# Generar documentos
# ---------------------------------
print("\nGenerando documentos Word...\n")

for index, fila in df.iterrows():

    doc = Document(template_path)

    reemplazos = {}

    for columna in df.columns:

        valor = fila[columna]

        reemplazos[str(columna)] = (
            ""
            if pd.isna(valor)
            else valor
        )

    for p in doc.paragraphs:

        reemplazar_en_parrafo(
            p,
            reemplazos
        )

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for p in cell.paragraphs:

                    reemplazar_en_parrafo(
                        p,
                        reemplazos
                    )

    partes_nombre = []

    for col in columnas_nombre:

        valor = str(
            fila[col]
        ).strip()

        partes_nombre.append(
            valor
        )

    if texto_adicional:

        partes_nombre.insert(
            0,
            texto_adicional
        )

    nombre_archivo = " ".join(
        partes_nombre
    )

    caracteres_invalidos = r'\/:*?"<>|'

    for c in caracteres_invalidos:

        nombre_archivo = nombre_archivo.replace(
            c,
            ""
        )

    nombre_archivo += ".docx"

    ruta_salida = os.path.join(
        carpeta_generados,
        nombre_archivo
    )

    doc.save(ruta_salida)

    df.at[index, "Adjuntos"] = nombre_archivo

    print(
        f"Generado: {filaPrint} | {nombre_archivo}"
    )

    filaPrint += 1

# ---------------------------------
# Guardar Excel actualizado
# ---------------------------------
print("\nActualizando Excel...\n")

df.to_excel(
    excel_path,
    index=False
)

print("\nProceso de generación de Words terminado")

# ---------------------------------
# Conversión PDF
# ---------------------------------
print("\nIniciando conversión de Word a PDF...\n")

carpeta_origen = carpeta_generados

carpeta_destino_base = os.path.dirname(
    carpeta_generados
)

carpeta_origen = os.path.abspath(
    carpeta_origen
)

carpeta_destino_base = os.path.abspath(
    carpeta_destino_base
)

nombre_base = (
    os.path.basename(carpeta_origen)
    + " pdf"
)

carpeta_pdf = os.path.join(
    carpeta_destino_base,
    nombre_base
)

contador = 1
base_original = carpeta_pdf

while os.path.exists(carpeta_pdf):

    carpeta_pdf = (
        f"{base_original} ({contador})"
    )

    contador += 1

os.makedirs(carpeta_pdf)

word = win32com.client.DispatchEx(
    "Word.Application"
)

word.Visible = False

wdFormatPDF = 17

archivos = [
    f for f in os.listdir(carpeta_origen)
    if f.lower().endswith(".docx")
]

total = len(archivos)

for i, archivo in enumerate(archivos, start=1):

    ruta_docx = os.path.abspath(
        os.path.join(carpeta_origen, archivo)
    )

    nombre_pdf = (
        os.path.splitext(archivo)[0]
        + ".pdf"
    )

    ruta_pdf = os.path.abspath(
        os.path.join(carpeta_pdf, nombre_pdf)
    )

    try:

        doc = word.Documents.Open(
            ruta_docx,
            ReadOnly=True
        )

        doc.SaveAs(
            ruta_pdf,
            FileFormat=wdFormatPDF
        )

        doc.Close(False)

        print(
            f"[{i}/{total}] Convertido: {nombre_pdf}"
        )

    except Exception as e:

        print(
            f"\nError en:\n{ruta_docx}\n"
        )

        print(e)

# ---------------------------------
# Cerrar Word
# ---------------------------------
word.Quit()

print("\n==============================")
print(" PROCESO TERMINADO ")
print("==============================")

print(f"\nWords guardados en:\n{carpeta_generados}")
print(f"\nPDFs guardados en:\n{carpeta_pdf}")